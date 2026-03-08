import os
from PySide6.QtGui import QTextDocument, QFont, QColor, QPainter, QPageLayout, QPageSize
from PySide6.QtCore import QMarginsF
from PySide6.QtPrintSupport import QPrinter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Exporter:
    """Utility class for exporting content to PDF and DOCX."""

    @staticmethod
    def export_to_pdf(content_html: str, output_path: str, options: dict = None, resources: dict = None):
        """
        Exports HTML content to a PDF file.
        
        Args:
            content_html: The HTML content to export.
            output_path: The file path to save the PDF to.
            options: A dictionary of export options (font_size, line_spacing, margins, orientation, etc.)
            resources: A dictionary of {url: QImage} to insert into the document context natively.
        """
        if options is None:
            options = {}

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(output_path)

        # Set orientation
        orientation = options.get("orientation", "Portrait")
        if orientation == "Landscape":
            printer.setPageOrientation(QPageLayout.Orientation.Landscape)
        else:
            printer.setPageOrientation(QPageLayout.Orientation.Portrait)

        # Set margins (converted from mm to points or using QPageLayout)
        margins = options.get("margins", {"top": 10, "bottom": 10, "left": 10, "right": 10})
        layout = printer.pageLayout()
        layout.setMargins(QMarginsF(margins["left"], margins["top"], margins["right"], margins["bottom"]))
        printer.setPageLayout(layout)

        doc = QTextDocument()
        
        # Apply basic styling to the HTML for "Dark to Light" conversion
        font_size = options.get("font_size", 12)
        base_style = f"""
            <style>
                body {{
                    font-family: {options.get("font_family", "Times New Roman")};
                    font_size: {font_size}pt;
                    line-height: {options.get("line_spacing", 1.2)};
                    color: black;
                    background-color: white;
                }}
                h1 {{ font-size: {font_size * 2}pt; text-align: center; }}
                h2 {{ font-size: {font_size * 1.5}pt; }}
                .verse-number {{ color: gray; font-size: 0.8em; vertical-align: super; }}
            </style>
        """
        doc.setHtml(f"<html><head>{base_style}</head><body>{content_html}</body></html>")
        
        if resources:
            from PySide6.QtCore import QUrl
            for res_id, res_img in resources.items():
                doc.addResource(QTextDocument.ImageResource, QUrl(res_id), res_img)
                
        doc.print_(printer)

    @staticmethod
    def export_to_docx(content_html: str, output_path: str, options: dict = None):
        """
        Exports HTML content to a DOCX file by parsing it natively using QTextDocument.
        
        Args:
            content_html: The HTML string to export.
            output_path: The file path to save the DOCX to.
            options: A dictionary of export options.
        """
        if options is None:
            options = {}

        doc = Document()
        
        font_size = options.get("font_size", 12)
        font_family = options.get("font_family", "Times New Roman")

        # Set default font
        style = doc.styles['Normal']
        style.font.name = font_family
        style.font.size = Pt(font_size)

        qt_doc = QTextDocument()
        
        base_style = f"""
            <style>
                body {{
                    font-family: {font_family};
                    font_size: {font_size}pt;
                    color: black;
                }}
                h1 {{ font-size: {font_size * 2}pt; text-align: center; }}
                h2 {{ font-size: {font_size * 1.5}pt; }}
            </style>
        """
        qt_doc.setHtml(f"<html><head>{base_style}</head><body>{content_html}</body></html>")
        
        from PySide6.QtGui import QTextListFormat
        from PySide6.QtCore import Qt

        block = qt_doc.begin()
        while block.isValid():
            block_fmt = block.blockFormat()
            lst = block.textList()
            if lst:
                lst_fmt = lst.format()
                indent = lst_fmt.indent()
                style_num = lst_fmt.style()

                is_bullet = style_num in (QTextListFormat.ListDisc, QTextListFormat.ListCircle, QTextListFormat.ListSquare)

                if is_bullet:
                    style_name = 'List Bullet'
                    if indent > 1:
                        style_name = f'List Bullet {min(indent, 3)}'
                else:
                    style_name = 'List Number'
                    if indent > 1:
                        style_name = f'List Number {min(indent, 3)}'

                try:
                    p = doc.add_paragraph(style=style_name)
                except KeyError:
                    p = doc.add_paragraph()
            else:
                heading_level = block_fmt.headingLevel()
                if heading_level > 0:
                    p = doc.add_paragraph(style=f'Heading {min(heading_level, 9)}')
                    # Ensure heading is black
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    p = doc.add_paragraph()
                    
            # Alignment
            align = block_fmt.alignment()
            if align & Qt.AlignLeft:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif align & Qt.AlignRight:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align & Qt.AlignHCenter:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align & Qt.AlignJustify:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            p.paragraph_format.line_spacing = options.get("line_spacing", 1.2)
            
            # Indentation
            # Qt 'indent' property is a multiplier of a standard indentation (usually 40px)
            # Qt 'leftMargin' is absolute pixels.
            qt_indent = block_fmt.indent()
            qt_margin = block_fmt.leftMargin()
            if qt_indent > 0 or qt_margin > 0:
                # Approximate 72dpi: 1 pixel = 1 point
                total_indent_pts = (qt_indent * 40) + qt_margin
                p.paragraph_format.left_indent = Pt(total_indent_pts)
            
            it = block.begin()
            while not it.atEnd():
                fragment = it.fragment()
                if fragment.isValid():
                    frag_text = fragment.text()
                    # Filter out Qt block separator chars
                    frag_text = frag_text.replace('\u2028', '\n').replace('\u2029', '\n')
                    
                    if frag_text:
                        run = p.add_run(frag_text)
                        char_fmt = fragment.charFormat()
                        
                        if char_fmt.fontWeight() == QFont.Bold or char_fmt.fontWeight() == QFont.Weight.Bold:
                            run.bold = True
                        if char_fmt.fontItalic():
                            run.italic = True
                        if char_fmt.fontUnderline():
                            run.underline = True
                        if char_fmt.fontStrikeOut():
                            run.font.strike = True
                            
                        pt_size = char_fmt.fontPointSize()
                        if pt_size > 0 and pt_size != font_size:
                            run.font.size = Pt(pt_size)
                            
                        from PySide6.QtGui import QTextFormat
                        if char_fmt.hasProperty(QTextFormat.ForegroundBrush):
                            fg_color = char_fmt.foreground().color()
                            if fg_color.isValid() and fg_color.name() != '#000000':
                                run.font.color.rgb = RGBColor(fg_color.red(), fg_color.green(), fg_color.blue())
                                
                        if char_fmt.hasProperty(QTextFormat.BackgroundBrush):
                            bg_color = char_fmt.background().color()
                            if bg_color.isValid():
                                from docx.enum.text import WD_COLOR_INDEX
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                
                        if char_fmt.isAnchor():
                            run.underline = True
                            if not char_fmt.hasProperty(QTextFormat.ForegroundBrush):
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
                it += 1
                
            block = block.next()

        doc.save(output_path)
